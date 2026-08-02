"""Render an InventoryReport to a Word document (.docx).

Design rules:
  - Beautiful, industry-style layout: cover page, table of contents,
    sectioned body, branded header/footer, readable tables.
  - Pure python-docx, no system deps.
  - Everything bounded by the same caps the collector uses; render is
    a straight walk of the already-bounded report.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from fpulse.reports.inventory import InventoryReport, ProjectInventory, PipelineInventory
from fpulse.reports.cron_human import cron_to_human


# Brand palette — matches F-Pulse UI accents.
BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)        # slate-900
BRAND_VIOLET = RGBColor(0x7C, 0x3A, 0xED)      # violet-600
BRAND_MUTED = RGBColor(0x64, 0x74, 0x8B)       # slate-500
BRAND_BG = RGBColor(0xF8, 0xFA, 0xFC)          # slate-50
STATUS_GREEN = RGBColor(0x05, 0x96, 0x69)
STATUS_AMBER = RGBColor(0xD9, 0x77, 0x06)
STATUS_RED = RGBColor(0xDC, 0x26, 0x26)


def render_docx(report: InventoryReport) -> bytes:
    """Produce the complete .docx as a bytes blob.

    Section layout is tier-aware:
      plus — cover, TOC, exec summary, ops audit, projects, connections,
             users, schedules, alerts, approval gates, appendix.
      free — cover, TOC, exec summary, ops audit (compact), projects,
             connections, schedules, alerts, upgrade CTA, appendix.
             (Users and Approval Gates are Plus-only sections.)
    """
    is_free = report.tier == "free"
    doc = Document()
    _configure_styles(doc)
    _add_cover_page(doc, report)
    _add_toc(doc)
    # "What should I fix first?" headline — appears BEFORE the
    # executive summary so the reader sees actionable signal first.
    _add_insights_section(doc, report)
    # 2026-06-05 — Steward reliability snapshot. Plain-text equivalent
    # of the PDF version. Skipped if Steward is disabled or has no
    # findings.
    _add_steward_section(doc, report)
    _add_executive_summary(doc, report)
    _add_operational_audit_section(doc, report)
    _add_failure_analysis_section(doc, report)
    _add_duration_analysis_section(doc, report)
    _add_projects_section(doc, report)
    _add_connections_section(doc, report)
    if not is_free:
        _add_users_section(doc, report)
    _add_schedules_section(doc, report)
    _add_alerts_section(doc, report)
    if not is_free:
        _add_approval_gates_section(doc, report)
    if is_free:
        _add_upgrade_cta(doc, report)
    _add_appendix(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# Styling
# ═══════════════════════════════════════════════════════════════════════


def _configure_styles(doc: Document) -> None:
    """Tighten default margins, install heading fonts."""
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    styles = doc.styles
    # Normal paragraph
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = BRAND_DARK

    for level, (size, color) in enumerate([
        (22, BRAND_DARK),       # Heading 1
        (16, BRAND_VIOLET),     # Heading 2
        (13, BRAND_DARK),       # Heading 3
        (11, BRAND_DARK),       # Heading 4
    ], start=1):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


# ═══════════════════════════════════════════════════════════════════════
# Cover page
# ═══════════════════════════════════════════════════════════════════════


def _add_cover_page(doc: Document, report: InventoryReport) -> None:
    is_free = report.tier == "free"
    product_name = "F-PULSE" if is_free else "F-PULSE"
    product_version_label = "F-Pulse version" if is_free else "F-Pulse version"

    # Title band
    for _ in range(3):
        doc.add_paragraph()  # vertical padding

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(product_name)
    run.font.size = Pt(36)
    run.font.bold = True
    # Free tier gets a calmer, slate color to distinguish tier at a glance.
    run.font.color.rgb = BRAND_MUTED if is_free else BRAND_VIOLET

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = "System Report" if is_free else "System Inventory Report"
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Workspace: {report.workspace_name}")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_DARK

    # Environment filter tag — prominent when set.
    if report.env_filter != "all":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Environment: {report.env_filter.upper()} only")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = (
            STATUS_RED if report.env_filter == "prod" else BRAND_VIOLET
        )

    doc.add_paragraph()

    # Metadata block
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ("Generated at",  _fmt_iso(report.generated_at)),
        ("Generated by",  report.generated_by),
        ("Report scope",
         "Workspace (full)" if report.tier == "free"
         else ("Administrator (full workspace)" if report.scope == "admin"
               else "User (ACL-filtered)")),
        (product_version_label,
         f"{report.fpulse_version} (schema v{report.schema_version})"),
    ]):
        row = meta_tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        _style_meta_row(row.cells[0], bold=True, muted=True)
        _style_meta_row(row.cells[1])

    # Bottom caption
    for _ in range(4):
        doc.add_paragraph()

    # Redaction notice — trust signal that this document is safe to share.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "🔒  All credentials in this report are redacted. "
        "Secrets are shown as Vault references or marked [INLINE — MIGRATE]. "
        "This document is safe to share via email, ticketing, or print."
    )
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = STATUS_GREEN

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This report describes the live state of your F-Pulse installation "
        "at the moment of generation. For the current state, regenerate from "
        "the Reports page.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_MUTED

    _page_break(doc)


def _style_meta_row(cell, *, bold: bool = False, muted: bool = False) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10.5)
            if bold:
                run.font.bold = True
            if muted:
                run.font.color.rgb = BRAND_MUTED
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_schedule_cell(cell, cron_expression: str) -> None:
    """Write a humanized cron description + raw expression into a cell.

    Two stacked lines:
      1. "Every 2 minutes"  — regular weight, default color
      2. "*/2 * * * *"     — smaller, muted (slate-500)
    If the cron pattern isn't one we recognize, we only write the raw
    expression so we don't duplicate it.
    """
    human = cron_to_human(cron_expression)
    # Clear existing default-empty paragraph, then write our own.
    first = cell.paragraphs[0]
    first.text = ""
    first.add_run(human)
    if human != cron_expression and cron_expression:
        sub = cell.add_paragraph()
        run = sub.add_run(cron_expression)
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_MUTED


# ═══════════════════════════════════════════════════════════════════════
# Table of contents
# ═══════════════════════════════════════════════════════════════════════


def _add_toc(doc: Document) -> None:
    """Insert a Word field-based TOC. Word will populate it on first
    open (right-click → Update Field). This is the standard approach."""
    h = doc.add_heading("Table of Contents", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # field begin
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    # field code
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    # separator
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    # placeholder text (shown until Word updates the field)
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click → Update Field to populate the table of contents."
    # field end
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    rElement = run._r
    rElement.append(fldChar1)
    rElement.append(instrText)
    rElement.append(fldChar2)
    rElement.append(fldChar3)
    rElement.append(fldChar4)

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Executive summary
# ═══════════════════════════════════════════════════════════════════════


def _add_executive_summary(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        f"This report describes the live state of workspace "
        f"{report.workspace_name!r} in the F-Pulse installation at "
        f"{_fmt_iso(report.generated_at)}. All counts and lists are "
        "sourced directly from the backing stores at the time of "
        "generation."
    )

    # Totals table
    totals = report.totals
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Count"

    is_free = report.tier == "free"
    rows = [
        ("Projects", totals.get("projects", 0)),
        ("Pipelines (total)", totals.get("pipelines", 0)),
        ("  ↳ DEV only (not in PROD)", totals.get("pipelines_dev_only", 0)),
        ("  ↳ Live in PROD (deployed)", totals.get("pipelines_in_prod", 0)),
        ("Connections (total)", totals.get("connections", 0)),
        ("  ↳ DEV-scoped",     totals.get("connections_dev", 0)),
        ("  ↳ PROD-scoped",    totals.get("connections_prod", 0)),
        ("  ↳ Shared (unset)", totals.get("connections_shared", 0)),
        ("  ↳ With inline credentials (migrate candidates)",
         totals.get("connections_inline_creds", 0)),
    ]
    # Users + Approval Gates are Plus-only concepts — OSS Free has a
    # single bootstrap user and no approvals surface.
    if not is_free:
        rows.append(("Users (active / total)",
                     f"{totals.get('users_active', 0)} / {totals.get('users', 0)}"))
    rows.extend([
        ("Schedules (enabled / total)",
         f"{totals.get('schedules_enabled', 0)} / {totals.get('schedules', 0)}"),
        ("Alert rules (enabled / total)",
         f"{totals.get('alerts_enabled', 0)} / {totals.get('alerts', 0)}"),
    ])
    if not is_free:
        rows.append(("Approval gates", totals.get("approval_gates", 0)))
    for label, val in rows:
        row = tbl.add_row().cells
        row[0].text = str(label)
        row[1].text = str(val)

    # Health
    doc.add_paragraph()
    doc.add_heading("1.1 Installation Health", level=2)

    p = doc.add_paragraph()
    run = p.add_run(f"Health score: {report.health.get('score', 0)} / 100")
    run.font.bold = True
    run.font.size = Pt(12)
    score = report.health.get("score", 0)
    if score >= 80:
        run.font.color.rgb = STATUS_GREEN
    elif score >= 50:
        run.font.color.rgb = STATUS_AMBER
    else:
        run.font.color.rgb = STATUS_RED

    issues = report.health.get("issues", [])
    if not issues:
        p = doc.add_paragraph("No issues detected — system is in clean shape.")
        p.runs[0].font.color.rgb = STATUS_GREEN
    else:
        doc.add_paragraph("Issues requiring attention:")
        for issue in issues:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(issue)
            run.font.size = Pt(10.5)

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Operational Audit (section 2)
# ═══════════════════════════════════════════════════════════════════════


def _add_insights_section(doc: Document, report: InventoryReport) -> None:
    """DOCX twin of the PDF insights section — same data, plain layout."""
    insights = report.insights
    if not insights or (insights.failing_count + insights.stale_count + insights.healthy_count) == 0:
        return
    doc.add_heading("What should I fix first?", level=1)
    doc.add_paragraph(insights.top_action)
    summary = doc.add_paragraph()
    summary.add_run(f"Failing: {insights.failing_count}    ").bold = True
    summary.add_run(f"Stale: {insights.stale_count}    ").bold = True
    summary.add_run(f"Healthy: {insights.healthy_count}").bold = True
    for item in insights.items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item.icon} {item.headline}")
        run.bold = True
        if item.detail:
            p.add_run(f" — {item.detail}")
    doc.add_paragraph()  # trailing blank


def _add_steward_section(doc: Document, report: InventoryReport) -> None:
    """DOCX twin of the PDF Steward snapshot. Same data, plain layout."""
    sw = report.steward_summary or {}
    if not sw.get("enabled"):
        return
    total = sw.get("total_open_findings", 0)
    if total == 0:
        return
    by_sev = sw.get("by_severity") or {}
    by_kind = sw.get("by_kind") or {}
    mem = sw.get("memory_stats") or {}

    doc.add_heading("Steward — reliability snapshot", level=1)
    doc.add_paragraph(
        f"The F-Pulse Steward (read-only reliability + learning layer) sees "
        f"{total} open finding{'s' if total != 1 else ''} in this workspace "
        f"at report time. Open the eye icon in the app header for full detail."
    )
    summary = doc.add_paragraph()
    summary.add_run(f"P1 (escalated): {by_sev.get('p1', 0)}    ").bold = True
    summary.add_run(f"P2 (review): {by_sev.get('p2', 0)}    ").bold = True
    summary.add_run(f"P3 (info): {by_sev.get('p3', 0)}").bold = True

    kind_labels = {
        "duplicate_source": "Duplicate source",
        "duplicate_pipeline": "Duplicate pipeline",
        "failure_rca": "Failure RCA",
        "volume_anomaly": "Volume anomaly",
        "schema_drift": "Schema drift",
        "cost_recommendation": "Cost recommendation",
    }
    for kind, count in sorted(by_kind.items(), key=lambda kv: -kv[1]):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{kind_labels.get(kind, kind)}: ")
        run.bold = True
        p.add_run(str(count))

    scans = mem.get("total_scans", 0)
    dismisses = mem.get("total_dismisses", 0)
    resolves = mem.get("total_resolves", 0)
    if scans or dismisses or resolves:
        learn = doc.add_paragraph()
        run = learn.add_run(
            f"Learning history: {scans} scans recorded · "
            f"{dismisses} dismissed (intentional) · "
            f"{resolves} resolved by user action."
        )
        run.italic = True
    doc.add_paragraph()  # trailing blank


def _add_duration_analysis_section(doc: Document, report: InventoryReport) -> None:
    """Per-pipeline avg / p95 / last run + regression flag."""
    da = report.duration_analysis
    doc.add_heading("Run duration analysis (last 30 days)", level=1)
    if not da or da.total_pipelines == 0:
        doc.add_paragraph(
            "No pipelines with ≥3 runs in the last 30 days yet — duration "
            "analysis needs a minimum sample size to be meaningful."
        )
        return
    summary = doc.add_paragraph()
    summary.add_run(f"{da.total_pipelines} pipelines, {da.total_runs} runs total. ")
    if da.slowest_pipeline:
        summary.add_run(f"Slowest by p95: ")
        summary.add_run(da.slowest_pipeline).bold = True
        summary.add_run(". ")
    if da.regressions_count:
        run = summary.add_run(f"{da.regressions_count} regression suspect"
                              f"{'s' if da.regressions_count != 1 else ''}.")
        run.bold = True

    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Pipeline"
    hdr[1].text = "Runs"
    hdr[2].text = "Avg ms"
    hdr[3].text = "p95 ms"
    hdr[4].text = "Last ms"
    hdr[5].text = "Note"
    for r in da.rows[:15]:
        row = t.add_row().cells
        row[0].text = r.pipeline_name
        row[1].text = str(r.runs)
        row[2].text = f"{r.avg_ms:,}"
        row[3].text = f"{r.p95_ms:,}"
        row[4].text = f"{r.last_ms:,}"
        row[5].text = "regression" if r.regression else "—"
    doc.add_paragraph()


def _add_failure_analysis_section(doc: Document, report: InventoryReport) -> None:
    """30-day failure analysis — top failing pipelines + common errors."""
    fa = report.failure_analysis
    doc.add_heading("Failure analysis (last 30 days)", level=1)
    if not fa or fa.total_failures == 0:
        doc.add_paragraph("No failures in the last 30 days — pipelines are running clean.")
        return
    p = doc.add_paragraph()
    p.add_run(f"{fa.total_failures}").bold = True
    p.add_run(f" failures across ")
    p.add_run(f"{fa.unique_failing_pipelines}").bold = True
    p.add_run(f" pipeline{'s' if fa.unique_failing_pipelines != 1 else ''} in the last {fa.window_days} days.")

    if fa.top_failing:
        doc.add_heading("Top failing pipelines", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Pipeline", "Failures", "Rate", "Last failure")
        for r in fa.top_failing[:10]:
            row = t.add_row().cells
            row[0].text = r.pipeline_name
            row[1].text = str(r.failure_count)
            row[2].text = f"{r.failure_rate_pct:.0f}% of {r.total_runs}" if r.total_runs else "—"
            row[3].text = r.last_failure_at or "—"

    if fa.top_errors:
        doc.add_heading("Most common error patterns", level=2)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = ("Error signature", "Count", "Affected pipelines")
        for r in fa.top_errors[:8]:
            row = t.add_row().cells
            row[0].text = r.error_signature[:120]
            row[1].text = str(r.count)
            affected = ", ".join(r.affected_pipelines[:4])
            if len(r.affected_pipelines) > 4:
                affected += f" + {len(r.affected_pipelines) - 4} more"
            row[2].text = affected or "—"

    doc.add_paragraph()


def _add_operational_audit_section(doc: Document, report: InventoryReport) -> None:
    audit = report.operational_audit
    doc.add_heading("2. Operational Audit", level=1)
    doc.add_paragraph(
        "Day-to-day operational signals: execution health over the last "
        f"{audit.window_hours} hours, pipelines that failed last run, and "
        "the next scheduled firings. This section is intended for data-ops "
        "teams checking system status at a glance."
    )

    # 2.1 Execution health (24h)
    doc.add_heading("2.1 Execution health (last 24 hours)", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 5"
    tbl.rows[0].cells[0].text = "Metric"
    tbl.rows[0].cells[1].text = "Value"

    rate = audit.success_rate_pct
    rate_str = f"{rate:.1f} %" if audit.total_executions else "— (no runs)"
    health_rows = [
        ("Total executions", audit.total_executions),
        ("Successful", audit.successful_executions),
        ("Failed / timed out", audit.failed_executions),
        ("Success rate", rate_str),
        ("Average duration",
         f"{audit.avg_duration_ms} ms" if audit.avg_duration_ms else "—"),
    ]
    for k, v in health_rows:
        row = tbl.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # 2.2 Recent failures
    doc.add_heading("2.2 Pipelines whose last run failed", level=2)
    if not audit.recent_failures:
        p = doc.add_paragraph("No pipelines are currently in a failed state.")
        p.runs[0].font.color.rgb = STATUS_GREEN
    else:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid Accent 5"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Pipeline", "Failed at", "Error"]):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for r in para.runs:
                    r.font.bold = True
        for f in audit.recent_failures:
            row = tbl.add_row().cells
            row[0].text = f.workflow_name
            row[1].text = _fmt_iso(f.failed_at)
            row[2].text = (f.error or "—")[:120]

    # 2.3 Next scheduled runs
    doc.add_heading("2.3 Next scheduled runs", level=2)
    if not audit.next_runs:
        doc.add_paragraph("No upcoming scheduled runs.")
    else:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 5"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Pipeline", "Schedule", "Env", "Next fire (UTC)"]):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for r in para.runs:
                    r.font.bold = True
        for nr in audit.next_runs:
            row = tbl.add_row().cells
            row[0].text = nr.workflow_name
            # 2026-05-26 — humanize cron. Two-line cell: human description
            # primary, raw expression on a small muted secondary line.
            _set_schedule_cell(row[1], nr.cron_expression)
            row[2].text = nr.environment
            row[3].text = _fmt_iso(nr.next_fire_at)

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Projects & pipelines (section 3)
# ═══════════════════════════════════════════════════════════════════════


def _add_projects_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("3. Projects", level=1)

    if not report.projects:
        doc.add_paragraph("No projects visible at the report's scope.")
        _page_break(doc)
        return

    doc.add_paragraph(
        f"This installation has {len(report.projects)} project(s). Each "
        "project below lists its owner, approver, members, and the pipelines "
        "inside it with their connections, schedules, alerts, and recent runs."
    )

    for i, proj in enumerate(report.projects, start=1):
        _add_project_block(doc, proj, f"3.{i}")

    _page_break(doc)


def _add_project_block(doc: Document, proj: ProjectInventory, section: str) -> None:
    doc.add_heading(f"{section} {proj.name}", level=2)

    # Header table
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Shading Accent 5"
    _kv_row(tbl, "Project ID", proj.id)
    _kv_row(tbl, "Description", proj.description or "—")
    _kv_row(tbl, "Owner", proj.owner or "(unset)")
    _kv_row(tbl, "Approver", proj.approver or "(all admins)")
    _kv_row(tbl, "Approval status", proj.approval_status or "none")
    _kv_row(tbl, "Members",
            (", ".join(proj.member_names) if proj.member_names else "—"))
    _kv_row(tbl, "Created", _fmt_iso(proj.created_at))
    _kv_row(tbl, "Pipelines in project", str(proj.pipeline_count))

    doc.add_paragraph()

    if not proj.pipelines:
        p = doc.add_paragraph("No pipelines in this project yet.")
        p.runs[0].italic = True
        return

    doc.add_heading(f"{section}.1 Pipelines", level=3)
    for j, p_inv in enumerate(proj.pipelines, start=1):
        _add_pipeline_block(doc, p_inv, f"{section}.1.{j}")


def _add_pipeline_block(doc: Document, p: PipelineInventory, section: str) -> None:
    doc.add_heading(f"{section} {p.name}", level=4)

    # ── Prominent signals line (gap 4) ────────────────────────────────
    # Shows env badge(s) + last-run pill + next-run time BEFORE the full
    # property table, so ops teams can scan status at a glance.
    signals = doc.add_paragraph()
    signals.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Env badges
    for env in p.environments:
        run = signals.add_run(f"  {env}  ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = (
            STATUS_RED if env == "PROD" else BRAND_VIOLET
        )

    signals.add_run("   ")

    # Last-run pill
    lrs = (p.last_run_status or "").lower()
    if lrs == "success":
        run = signals.add_run("LAST RUN: SUCCESS")
        run.font.color.rgb = STATUS_GREEN
    elif lrs in ("error", "failed", "timeout"):
        run = signals.add_run(f"LAST RUN: {lrs.upper()}")
        run.font.color.rgb = STATUS_RED
    elif lrs:
        run = signals.add_run(f"LAST RUN: {lrs.upper()}")
        run.font.color.rgb = STATUS_AMBER
    else:
        run = signals.add_run("NEVER RUN")
        run.font.color.rgb = BRAND_MUTED
    run.font.bold = True
    run.font.size = Pt(9)

    if p.last_run_at:
        r2 = signals.add_run(f"  ({_fmt_iso(p.last_run_at)})")
        r2.font.size = Pt(8)
        r2.font.color.rgb = BRAND_MUTED

    if p.next_run_at:
        signals.add_run("   ")
        r3 = signals.add_run(f"NEXT: {_fmt_iso(p.next_run_at)}")
        r3.font.bold = True
        r3.font.size = Pt(9)
        r3.font.color.rgb = BRAND_DARK

    # ── Purpose (the "why") — prominent, right under the name ──────────
    if p.business_purpose:
        pp = doc.add_paragraph()
        pr = pp.add_run("Purpose: ")
        pr.font.bold = True
        pp.add_run(p.business_purpose)

    # ── Property table ────────────────────────────────────────────────
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 5"
    _kv_row(tbl, "Environment", " + ".join(p.environments))
    _kv_row(tbl, "Status", _status_label(p.status, p.deployed_version))
    _kv_row(tbl, "Latest version", f"v{p.latest_version}")
    _kv_row(tbl, "Deployed version",
            f"v{p.deployed_version}" if p.deployed_version else "not deployed")
    _kv_row(tbl, "Owner", p.owner or "—")
    if p.description:
        _kv_row(tbl, "Description", p.description)
    if p.readme:
        _kv_row(tbl, "Notes (README)", p.readme)
    if p.approval_status:
        _kv_row(tbl, "Approval",
                f"{p.approval_status}" +
                (f" by {p.approved_by}" if p.approved_by else "") +
                (f" — submitted by {p.submitted_by}" if p.submitted_by else ""))
    _kv_row(tbl, "Nodes", f"{p.step_count} steps ({len(p.node_types)} types)")
    if p.node_types:
        _kv_row(tbl, "Node types", ", ".join(p.node_types))

    if p.connections_used:
        conn_lines = [f"{c['name']} ({c['type']})" for c in p.connections_used]
        _kv_row(tbl, "Connections used", "; ".join(conn_lines))

    if p.schedules:
        sch_lines = []
        for s in p.schedules:
            enabled = "✓" if s.get("enabled") else "✗"
            sch_lines.append(f"{enabled} {cron_to_human(s.get('cron', '?'))} "
                             f"({s.get('timezone', 'UTC')}, "
                             f"{s.get('environment', 'DEV')})")
        _kv_row(tbl, "Schedules", "\n".join(sch_lines))
    else:
        _kv_row(tbl, "Schedules", "none")

    if p.alert_rules:
        alert_lines = []
        for a in p.alert_rules:
            enabled = "✓" if a.get("enabled") else "✗"
            channels = ", ".join(a.get("channels", [])) or "—"
            alert_lines.append(f"{enabled} {a.get('name', '(unnamed)')} "
                               f"on {a.get('condition', '?')} → {channels}")
        _kv_row(tbl, "Alert rules", "\n".join(alert_lines))
    else:
        _kv_row(tbl, "Alert rules", "none")

    if p.last_runs:
        run_lines = []
        for r in p.last_runs:
            run_lines.append(
                f"{r.get('status', '?')} | {r.get('duration_ms', 0)} ms | "
                f"{r.get('rows', 0)} rows | "
                f"{_fmt_iso(r.get('started_at', ''))} | "
                f"by {r.get('triggered_by', '—')}"
            )
        _kv_row(tbl, f"Last {len(run_lines)} runs", "\n".join(run_lines))
    else:
        _kv_row(tbl, "Recent runs", "no executions recorded")

    if p.content_hash:
        _kv_row(tbl, "Content hash",
                p.content_hash[:16] + "… (signed artifact)")

    doc.add_paragraph()  # breathing room


# ═══════════════════════════════════════════════════════════════════════
# Connections
# ═══════════════════════════════════════════════════════════════════════


def _add_connections_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("4. Connections", level=1)

    if not report.connections:
        doc.add_paragraph("No connections configured.")
        _page_break(doc)
        return

    doc.add_paragraph(
        f"This installation has {len(report.connections)} saved connection(s). "
        "Connections are reusable references to data sources and sinks; each "
        "holds non-secret config plus a pointer to credentials in the Vault."
    )

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Name", "Type", "Env", "Capabilities", "Creds", "Used by"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for c in report.connections:
        row = tbl.add_row().cells
        row[0].text = c.name
        row[1].text = c.type
        row[2].text = c.environment or "—"
        row[3].text = ", ".join(c.capabilities) or "—"
        # Redaction marker (gap 2) — show Vault reference masked, or a
        # clear "inline — migrate" warning. Never shows raw secret text.
        if c.has_credential_ref:
            row[4].text = (f"Vault: {c.credential_ref}" if c.credential_ref
                           else "Vault (redacted)")
        elif c.has_inline_creds:
            row[4].text = "[INLINE — MIGRATE]"
        else:
            row[4].text = "none"
        used = c.used_by_pipelines
        row[5].text = ", ".join(used[:3]) + (f" (+{len(used) - 3})" if len(used) > 3 else "") \
            if used else "(unused)"

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Users
# ═══════════════════════════════════════════════════════════════════════


def _add_users_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("5. Users", level=1)

    if not report.users:
        doc.add_paragraph("No users found.")
        _page_break(doc)
        return

    doc.add_paragraph(
        f"{len(report.users)} user account(s) exist in this installation. "
        "Roles determine what each user can do; see the Security & Compliance "
        "guide for the full role matrix."
    )

    # By-role rollup
    by_role: dict[str, int] = {}
    for u in report.users:
        by_role[u.role] = by_role.get(u.role, 0) + 1

    doc.add_heading("5.1 By role", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 5"
    tbl.rows[0].cells[0].text = "Role"
    tbl.rows[0].cells[1].text = "Count"
    for role, count in sorted(by_role.items()):
        row = tbl.add_row().cells
        row[0].text = role
        row[1].text = str(count)

    # Full list (compact)
    doc.add_heading("5.2 Roster", level=2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Name", "Email", "Role", "Active", "Last login"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for u in report.users:
        row = tbl.add_row().cells
        row[0].text = u.name or "—"
        row[1].text = u.email or "—"
        row[2].text = u.role
        row[3].text = "yes" if u.is_active else "NO"
        row[4].text = _fmt_iso(u.last_login_at) if u.last_login_at else "—"

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Schedules
# ═══════════════════════════════════════════════════════════════════════


def _add_schedules_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("6. Schedules", level=1)

    if not report.schedules:
        doc.add_paragraph("No schedules configured.")
        _page_break(doc)
        return

    doc.add_paragraph(
        f"{len(report.schedules)} schedule(s) configured. PROD schedules "
        "execute the pipeline's pinned deployed_version; DEV schedules execute "
        "the latest saved version."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Pipeline", "Schedule", "Timezone", "Env", "Next fire"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for s in report.schedules:
        row = tbl.add_row().cells
        row[0].text = s.workflow_name or s.workflow_id
        # 2026-05-26 — humanize cron in the Schedule cell.
        _set_schedule_cell(row[1], s.cron_expression)
        row[2].text = s.timezone
        row[3].text = s.environment
        row[4].text = _fmt_iso(s.next_fire_at) if s.next_fire_at else "—"

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Alerts
# ═══════════════════════════════════════════════════════════════════════


def _add_alerts_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("7. Alert Rules", level=1)

    if not report.alerts:
        doc.add_paragraph("No alert rules configured.")
        _page_break(doc)
        return

    doc.add_paragraph(
        f"{len(report.alerts)} alert rule(s) configured. Alerts fire on "
        "execution outcome or runtime metric and notify via the configured "
        "channels (in-app, email, Slack, webhook)."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Rule", "Pipeline", "Condition", "Channels", "Enabled"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for a in report.alerts:
        row = tbl.add_row().cells
        row[0].text = a.name
        row[1].text = a.workflow_name or a.workflow_id
        row[2].text = a.condition
        row[3].text = ", ".join(a.channels) or "—"
        row[4].text = "✓" if a.enabled else "✗"

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Approval gates
# ═══════════════════════════════════════════════════════════════════════


def _add_approval_gates_section(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("8. Approval Gates", level=1)

    if not report.approval_gates:
        doc.add_paragraph(
            "No approval gates configured. PROD deploys fall back to notifying "
            "all admins and leads in the workspace. For tighter governance, "
            "configure per-project or per-pipeline gates."
        )
        _page_break(doc)
        return

    doc.add_paragraph(
        f"{len(report.approval_gates)} approval gate(s) configured. Gates "
        "resolve most-specific-wins: pipeline → project → global."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 5"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Scope", "Scope ID", "Min approvals", "Approvers", "Channels"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for g in report.approval_gates:
        row = tbl.add_row().cells
        row[0].text = g.scope
        row[1].text = g.scope_id or "(global)"
        row[2].text = str(g.min_approvals)
        row[3].text = ", ".join(g.approvers) or "—"
        row[4].text = ", ".join(g.notify_channels) or "—"

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
# Appendix
# ═══════════════════════════════════════════════════════════════════════


def _add_upgrade_cta(doc: Document, report: InventoryReport) -> None:
    """Free-tier only — a friendly, non-spammy callout describing what
    F-Pulse adds on top of the report the reader just reviewed."""
    doc.add_heading("Upgrade to F-Pulse", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "The report you are reading covers the core F-Pulse inventory. "
        "F-Pulse extends this report with enterprise-grade sections that "
        "are not available on the free tier:"
    )
    run.font.size = Pt(11)

    upgrade_items = [
        ("User roster & role matrix",
         "Every user, their role, last login, and environment permissions."),
        ("Approval gates",
         "Per-pipeline, per-project, and global gates with approver lists."),
        ("Operational audit (full)",
         "24-hour success rate across every execution, with failure roster "
         "and the next 10 scheduled firings in one place."),
        ("Signed artifacts & audit log",
         "Tamper-evident pipeline deploys with SHA-256 content hashes and "
         "a full audit of every state-changing action."),
        ("Vault-backed credentials",
         "AES-256 encrypted secret storage with rotation, audit, and "
         "masked references in reports."),
        ("DEV → PROD approval workflow",
         "Submit, review, approve, deploy, rollback — all tracked."),
        ("Drift detection & retention",
         "Nightly schema-drift scans + two-tier retention with Parquet "
         "archive before purge."),
        ("Prometheus metrics, Grafana dashboards, Loki logs",
         "Production-grade observability, auto-provisioned."),
    ]
    for title, desc in upgrade_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title} — ")
        run.font.bold = True
        run.font.color.rgb = BRAND_VIOLET
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Contact your F-Pulse account representative or visit the Admin "
        "page inside the application to activate a Plus license."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_MUTED

    _page_break(doc)


def _add_appendix(doc: Document, report: InventoryReport) -> None:
    doc.add_heading("Appendix A — Report Metadata", level=1)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Shading Accent 5"
    _kv_row(tbl, "Generated at", _fmt_iso(report.generated_at))
    _kv_row(tbl, "Generated by", report.generated_by)
    _kv_row(tbl, "Scope", report.scope)
    _kv_row(tbl, "Tier", report.tier.upper())
    _kv_row(tbl, "Environment filter",
            "all environments" if report.env_filter == "all"
            else f"{report.env_filter.upper()} only")
    _kv_row(tbl, "Workspace ID", report.workspace_id)
    _kv_row(tbl, "Product version",
            f"{'F-Pulse' if report.tier == 'free' else 'F-Pulse'} "
            f"{report.fpulse_version}")
    _kv_row(tbl, "Schema version", f"v{report.schema_version}")
    _kv_row(tbl, "Report format", "Microsoft Word (.docx)")

    doc.add_paragraph()
    p = doc.add_paragraph(
        "This report was generated by F-Pulse from live data and reflects the "
        "state of the system at the timestamp above. Regenerate from the "
        "Reports page in the F-Pulse UI at any time."
    )
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = BRAND_MUTED


# ═══════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════


def _kv_row(tbl, key: str, value: str) -> None:
    row = tbl.add_row().cells
    row[0].text = key
    row[1].text = str(value) if value is not None else "—"
    # Bold + muted the key cell
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = BRAND_MUTED
            run.font.size = Pt(10)
    for para in row[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _fmt_iso(s: str) -> str:
    if not s:
        return "—"
    try:
        # Accept "...Z" and "...+00:00"
        if s.endswith("Z"):
            s_norm = s.replace("Z", "+00:00")
        else:
            s_norm = s
        dt = datetime.fromisoformat(s_norm)
        return dt.strftime("%Y-%m-%d %H:%M UTC")
    except (ValueError, TypeError):
        return s


def _status_label(status: str, deployed: int | None) -> str:
    base = (status or "draft").lower()
    if deployed:
        return f"{base.upper()} (PROD v{deployed})"
    return base.upper()
